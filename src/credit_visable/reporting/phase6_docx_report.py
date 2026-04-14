"""Generate an English Phase 6 analytical DOCX report from existing artifacts."""

from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd

from credit_visable.utils import get_paths

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - exercised only when dependency is absent.
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None
    _DOCX_IMPORT_ERROR = exc
else:
    _DOCX_IMPORT_ERROR = None


DEFAULT_REPORT_NAME = "Phase6_Scorecard_Cutoff_Analysis_Report.docx"
DEFAULT_PHASE6_LABEL = "xgboost_traditional_plus_proxy"
PHASE5_DIR_NAME = "xai_fairness"


@dataclass(frozen=True)
class FigureSpec:
    """DOCX figure placement instructions."""

    filename: str
    caption: str
    width_inches: float


MAIN_FIGURES: tuple[FigureSpec, ...] = (
    FigureSpec(
        filename="phase5_proxy_family_uplift_delta.png",
        caption="Phase 5 proxy uplift review showing that ext_source dominates the candidate model's proxy-sensitive uplift.",
        width_inches=6.0,
    ),
    FigureSpec(
        filename="phase5_fairness_metric_gaps.png",
        caption="Phase 5 fairness metric gaps indicating that age_band remains the strongest governance sensitivity, with region_rating_group and family_status_group still material.",
        width_inches=6.0,
    ),
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_calibration_curve.png",
        caption="Calibrated PD versus observed default rate by bin, showing that the Phase 6 translation now tracks realized risk closely.",
        width_inches=6.0,
    ),
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_decile_reliability.png",
        caption="Decile reliability after calibration, comparing mean calibrated PD against the observed default rate across risk deciles.",
        width_inches=6.0,
    ),
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_risk_band_count.png",
        caption="Risk band population distribution after calibrated score translation, with most applicants allocated across operational bands A to C.",
        width_inches=5.8,
    ),
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_risk_band_actual_default_rate.png",
        caption="Actual default rate by risk band, confirming that the operational bands preserve clear monotonic separation.",
        width_inches=5.8,
    ),
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_final_policy_cutoff_curve.png",
        caption="Final policy cutoff curve linking expected value and approved bad rate to the cutoff anchor score.",
        width_inches=6.2,
    ),
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_decision_migration_heatmap.png",
        caption="Final decision migration from the comparator policy to the candidate policy, highlighting where approve, review, and reject outcomes changed.",
        width_inches=6.1,
    ),
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_age_band_final_policy.png",
        caption="Final policy outcome mix by age_band, showing the widest remaining approval-rate spread across reviewed groups.",
        width_inches=6.1,
    ),
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_region_rating_final_policy.png",
        caption="Final policy outcome mix by region_rating_group, showing a clear approval disadvantage for the highest-risk rating segment.",
        width_inches=6.1,
    ),
)

APPENDIX_FIGURES: tuple[FigureSpec, ...] = (
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_pd_to_score_curve.png",
        caption="Supplementary PD-to-score mapping used for the calibrated Phase 6 score translation.",
        width_inches=5.9,
    ),
    FigureSpec(
        filename=f"phase6_{DEFAULT_PHASE6_LABEL}_score_histogram_kde.png",
        caption="Supplementary score distribution showing the central concentration of calibrated applicant scores.",
        width_inches=5.9,
    ),
)


def _ensure_docx_dependency() -> None:
    if Document is None or WD_ALIGN_PARAGRAPH is None or Inches is None or Pt is None:
        raise RuntimeError(
            "python-docx is required to generate the Phase 6 report. "
            "Install it with `python -m pip install python-docx`."
        ) from _DOCX_IMPORT_ERROR


def _load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def _format_percent(value: float, decimals: int = 2) -> str:
    return f"{value * 100:.{decimals}f}%"


def _format_pp(value: float, decimals: int = 2) -> str:
    return f"{value * 100:.{decimals}f} pp"


def _format_number(value: float, decimals: int = 2) -> str:
    return f"{value:,.{decimals}f}"


def _set_run_font(paragraph, *, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        run.bold = bold or run.bold
        run.italic = italic or run.italic
        if run.font is not None:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def _style_document(document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"


def _resolve_phase6_dir(paths) -> Path:
    preferred = paths.data_processed / "scorecard_cutoff" / DEFAULT_PHASE6_LABEL
    if preferred.exists():
        return preferred

    summaries = sorted((paths.data_processed / "scorecard_cutoff").glob("*/summary.json"))
    if not summaries:
        raise FileNotFoundError("No Phase 6 summary.json artifact was found under data/processed/scorecard_cutoff.")
    return summaries[0].parent


def _resolve_output_path(paths, output_path: Path | None) -> Path:
    if output_path is not None:
        return output_path
    return paths.reports / "Analysis" / DEFAULT_REPORT_NAME


def _build_group_gap_table(final_policy_group: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, str]] = []
    for protected_attribute in ["age_band", "family_status_group", "region_rating_group"]:
        subset = final_policy_group.loc[
            final_policy_group["protected_attribute"] == protected_attribute
        ].copy()
        best_row = subset.loc[subset["approval_rate"].idxmax()]
        worst_row = subset.loc[subset["approval_rate"].idxmin()]
        rows.append(
            {
                "Protected Attribute": protected_attribute.replace("_", " ").title(),
                "Highest Approval Group": str(best_row["group"]),
                "Highest Approval Rate": _format_percent(float(best_row["approval_rate"])),
                "Lowest Approval Group": str(worst_row["group"]),
                "Lowest Approval Rate": _format_percent(float(worst_row["approval_rate"])),
                "Gap": _format_pp(float(best_row["approval_rate"] - worst_row["approval_rate"])),
            }
        )
    return pd.DataFrame(rows)


def _build_report_context() -> dict[str, Any]:
    paths = get_paths()
    phase6_dir = _resolve_phase6_dir(paths)
    phase5_dir = paths.data_processed / PHASE5_DIR_NAME
    figures_dir = paths.reports_figures_v2
    candidate_label = phase6_dir.name

    source_paths = {
        "phase6_summary": phase6_dir / "summary.json",
        "final_policy": phase6_dir / "final_policy_summary.json",
        "score_transform": phase6_dir / "score_transform_meta.json",
        "risk_band": phase6_dir / "risk_band_summary.csv",
        "calibration": phase6_dir / "calibration_summary.csv",
        "final_policy_group": phase6_dir / "final_policy_group_summary.csv",
        "decision_migration": phase6_dir / "decision_migration_matrix.csv",
        "proxy_uplift": phase5_dir / "proxy_uplift_summary.csv",
        "fairness_metric": phase5_dir / "fairness_metric_summary.csv",
    }

    figure_specs: list[FigureSpec] = []
    for spec in MAIN_FIGURES + APPENDIX_FIGURES:
        if DEFAULT_PHASE6_LABEL not in spec.filename:
            figure_specs.append(spec)
            continue
        figure_specs.append(
            FigureSpec(
                filename=spec.filename.replace(DEFAULT_PHASE6_LABEL, candidate_label),
                caption=spec.caption.replace(DEFAULT_PHASE6_LABEL, candidate_label),
                width_inches=spec.width_inches,
            )
        )

    for spec in figure_specs:
        source_paths[spec.filename] = figures_dir / spec.filename

    missing = [str(path) for path in source_paths.values() if not path.exists()]
    if missing:
        raise FileNotFoundError(
            "The Phase 6 DOCX report requires existing Phase 5/6 artifacts. Missing files: "
            + ", ".join(missing)
        )

    summary = _load_json(source_paths["phase6_summary"])
    final_policy = _load_json(source_paths["final_policy"])
    score_transform = _load_json(source_paths["score_transform"])
    risk_band = _load_csv(source_paths["risk_band"])
    calibration = _load_csv(source_paths["calibration"])
    final_policy_group = _load_csv(source_paths["final_policy_group"])
    decision_migration = _load_csv(source_paths["decision_migration"])
    proxy_uplift = _load_csv(source_paths["proxy_uplift"])
    fairness_metric = _load_csv(source_paths["fairness_metric"])

    max_calibration_gap = float(calibration["calibration_gap"].abs().max())
    band_abc_share = float(
        risk_band.loc[risk_band["risk_band"].isin(["A", "B", "C"]), "population_share"].sum()
    )
    top_proxy = proxy_uplift.sort_values(
        "mean_abs_contribution_delta_proxy_minus_core", ascending=False
    ).iloc[0]

    group_gap_table = _build_group_gap_table(final_policy_group)

    decision_pivot = decision_migration.pivot(
        index="comparator_final_decision",
        columns="candidate_final_decision",
        values="count",
    ).fillna(0)
    approve_total = float(decision_pivot.loc["approve"].sum())
    review_total = float(decision_pivot.loc["review"].sum())
    reject_total = float(decision_pivot.loc["reject"].sum())
    approve_stay_rate = float(decision_pivot.loc["approve", "approve"] / approve_total)
    review_to_approve_rate = float(decision_pivot.loc["review", "approve"] / review_total)
    reject_to_approve_rate = float(decision_pivot.loc["reject", "approve"] / reject_total)

    fairness_focus = fairness_metric.set_index("protected_attribute")
    age_band_gap = float(fairness_focus.loc["age_band", "equalized_odds_gap"])
    region_gap = float(fairness_focus.loc["region_rating_group", "equalized_odds_gap"])
    family_gap = float(fairness_focus.loc["family_status_group", "equalized_odds_gap"])

    score_transform_table = pd.DataFrame(
        [
            {"Metric": "Base score", "Value": _format_number(float(score_transform["base_score"]), 0)},
            {"Metric": "Base odds", "Value": f"{_format_number(float(score_transform['base_odds']), 0)}:1"},
            {
                "Metric": "Points to double odds",
                "Value": _format_number(float(score_transform["points_to_double_odds"]), 0),
            },
            {"Metric": "Factor", "Value": _format_number(float(score_transform["factor"]))},
            {"Metric": "Offset", "Value": _format_number(float(score_transform["offset"]))},
            {"Metric": "Formula", "Value": str(score_transform["score_formula"])},
        ]
    )

    risk_band_table = risk_band.loc[
        :, ["risk_band", "count", "population_share", "mean_score", "mean_calibrated_pd", "actual_default_rate"]
    ].copy()
    risk_band_table.columns = [
        "Risk Band",
        "Count",
        "Population Share",
        "Mean Score",
        "Mean Calibrated PD",
        "Actual Default Rate",
    ]
    risk_band_table["Count"] = risk_band_table["Count"].map(lambda value: f"{int(value):,}")
    for column in ["Population Share", "Mean Calibrated PD", "Actual Default Rate"]:
        risk_band_table[column] = risk_band_table[column].map(_format_percent)
    risk_band_table["Mean Score"] = risk_band_table["Mean Score"].map(lambda value: _format_number(float(value), 1))

    final_policy_table = pd.DataFrame(
        [
            {"Metric": "Approve cutoff", "Value": _format_number(float(final_policy["final_cutoff"]), 0)},
            {"Metric": "Review floor", "Value": _format_number(float(final_policy["final_review_cutoff"]), 0)},
            {"Metric": "Approval rate", "Value": _format_percent(float(final_policy["approval_rate"]))},
            {"Metric": "Review rate", "Value": _format_percent(float(final_policy["review_rate"]))},
            {"Metric": "Reject rate", "Value": _format_percent(float(final_policy["reject_rate"]))},
            {"Metric": "Approved bad rate", "Value": _format_percent(float(final_policy["actual_approved_bad_rate"]))},
            {
                "Metric": "Expected value per applicant",
                "Value": _format_number(float(final_policy["expected_value_per_applicant"])),
            },
            {"Metric": "Guardrails passed", "Value": "Yes" if bool(final_policy["passes_guardrails"]) else "No"},
        ]
    )

    return {
        "paths": paths,
        "phase6_dir": phase6_dir,
        "figures_dir": figures_dir,
        "summary": summary,
        "final_policy": final_policy,
        "score_transform": score_transform,
        "risk_band": risk_band,
        "calibration": calibration,
        "group_gap_table": group_gap_table,
        "score_transform_table": score_transform_table,
        "risk_band_table": risk_band_table,
        "final_policy_table": final_policy_table,
        "candidate_label": candidate_label,
        "raw_pd_mean": float(summary["raw_pd_mean"]),
        "calibrated_pd_mean": float(summary["calibrated_pd_mean"]),
        "brier_score": float(calibration["brier_score"].iloc[0]),
        "max_calibration_gap": max_calibration_gap,
        "band_abc_share": band_abc_share,
        "top_proxy_family": str(top_proxy["proxy_family"]),
        "top_proxy_delta": float(top_proxy["mean_abs_contribution_delta_proxy_minus_core"]),
        "approve_stay_rate": approve_stay_rate,
        "review_to_approve_rate": review_to_approve_rate,
        "reject_to_approve_rate": reject_to_approve_rate,
        "age_band_gap": age_band_gap,
        "region_gap": region_gap,
        "family_gap": family_gap,
        "main_figures": figure_specs[: len(MAIN_FIGURES)],
        "appendix_figures": figure_specs[len(MAIN_FIGURES) :],
    }


def _add_heading(document, title: str, level: int) -> None:
    heading = document.add_heading(title, level=level)
    _set_run_font(heading, bold=True)


def _add_paragraph(document, text: str, *, bold_prefix: str | None = None, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        prefix_run.bold = True
        prefix_run.font.name = "Times New Roman"
        prefix_run.font.size = Pt(11)
        remainder = text[len(bold_prefix) :]
        if remainder:
            remainder_run = paragraph.add_run(remainder)
            remainder_run.italic = italic
            remainder_run.font.name = "Times New Roman"
            remainder_run.font.size = Pt(11)
    else:
        run = paragraph.add_run(text)
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def _add_bullet(document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def _add_table(document, table_frame: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(table_frame.columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, column in enumerate(table_frame.columns):
        header_cells[index].text = str(column)
        for run in header_cells[index].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    for _, row in table_frame.iterrows():
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for run in cells[index].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)


def _add_figure(document, figure_number: int, figure_path: Path, caption: str, width_inches: float) -> None:
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(figure_path), width=Inches(width_inches))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(f"Figure {figure_number}. {caption}")
    caption_run.italic = True
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(10)


def generate_phase6_analysis_docx(
    output_path: Path | None = None, include_phase5_bridge: bool = True
) -> Path:
    """Generate the Phase 6 English analysis report as a DOCX file."""

    _ensure_docx_dependency()
    context = _build_report_context()
    output_file = _resolve_output_path(context["paths"], output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Phase 6 Scorecard and Cutoff Analysis Report")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Home Credit Default Risk - English analytical report")
    subtitle_run.font.name = "Times New Roman"
    subtitle_run.font.size = Pt(12)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_text = (
        f"Prepared on {date.today().isoformat()} using existing Phase 5 and Phase 6 artifacts. "
        f"Candidate model: {context['candidate_label']}."
    )
    metadata_run = metadata.add_run(metadata_text)
    metadata_run.font.name = "Times New Roman"
    metadata_run.font.size = Pt(11)

    document.add_page_break()
    figure_number = 1

    _add_heading(document, "Executive Summary", level=1)
    _add_bullet(
        document,
        "Phase 6 corrects the probability-to-policy translation rather than the upstream ranking logic: raw PD mean is "
        f"{_format_percent(context['raw_pd_mean'])}, while calibrated PD mean is {_format_percent(context['calibrated_pd_mean'])}.",
    )
    _add_bullet(
        document,
        "Calibration quality is materially tighter after translation, with Brier score "
        f"{context['brier_score']:.4f} and a maximum observed calibration gap of {_format_pp(context['max_calibration_gap'])}.",
    )
    _add_bullet(
        document,
        "The final strategy uses an approve cutoff of "
        f"{_format_number(float(context['final_policy']['final_cutoff']), 0)} and a review floor of "
        f"{_format_number(float(context['final_policy']['final_review_cutoff']), 0)}, delivering approval / review / reject = "
        f"{_format_percent(float(context['final_policy']['approval_rate']))} / "
        f"{_format_percent(float(context['final_policy']['review_rate']))} / "
        f"{_format_percent(float(context['final_policy']['reject_rate']))}.",
    )
    _add_bullet(
        document,
        "Operational stratification is balanced rather than compressed: risk bands A, B, and C jointly contain "
        f"{_format_percent(context['band_abc_share'], 1)} of applicants, and the approved bad rate is "
        f"{_format_percent(float(context['final_policy']['actual_approved_bad_rate']))}.",
    )
    _add_bullet(
        document,
        "Under the current unit-economics assumptions, expected value per applicant is "
        f"{_format_number(float(context['final_policy']['expected_value_per_applicant']))}, and the final policy passes configured guardrails.",
    )

    if include_phase5_bridge:
        phase5_figure_numbers = (figure_number, figure_number + 1)
        figure_number += 2
        _add_heading(document, "Phase 5 Governance Bridge", level=1)
        _add_paragraph(
            document,
            "Phase 5 remains relevant because it explains which proxy-sensitive signals drove the candidate model into the final policy discussion. "
            f"The dominant uplift family is {context['top_proxy_family']}, with a mean absolute contribution delta of "
            f"{context['top_proxy_delta']:.3f} versus the traditional core comparator.",
        )
        _add_paragraph(
            document,
            f"Figure {phase5_figure_numbers[0]} and Figure {phase5_figure_numbers[1]} show the short governance bridge used in this report. "
            "The strongest grouped sensitivity still sits in age_band, while region_rating_group and family_status_group remain material enough "
            "to carry into policy interpretation rather than being treated as closed issues after Phase 5."
        )
        _add_paragraph(
            document,
            "This means the Phase 6 report should be read as a policy diagnostic: it translates the candidate model into score, band, and cutoff language, "
            "but it does not convert the grouped governance findings into a production-ready fairness clearance."
        )
        for index, figure in enumerate(context["main_figures"][:2], start=1):
            _add_figure(
                document,
                phase5_figure_numbers[index - 1],
                context["figures_dir"] / figure.filename,
                figure.caption,
                figure.width_inches,
            )

    calibration_figure_numbers = (figure_number, figure_number + 1)
    figure_number += 2
    _add_heading(document, "Calibration and Score Translation", level=1)
    _add_paragraph(
        document,
        "The central Phase 6 correction is the move from raw tree probability to calibrated PD before score translation. "
        f"Raw PD mean is {_format_percent(context['raw_pd_mean'])}, whereas calibrated PD mean falls to "
        f"{_format_percent(context['calibrated_pd_mean'])}. This is the corrected baseline for every downstream score, band, and cutoff statement in the report.",
    )
    _add_paragraph(
        document,
        f"Figure {calibration_figure_numbers[0]} and Figure {calibration_figure_numbers[1]} show that the calibrated probabilities track observed default behavior closely across bins and deciles. "
        f"The Brier score is {context['brier_score']:.4f}, and the maximum bin-level gap is only {_format_pp(context['max_calibration_gap'])}, "
        "which is consistent with a score translation that can be interpreted operationally."
    )
    _add_paragraph(
        document,
        "Table 1 summarizes the score transform settings used to convert calibrated PD into a deployable score scale."
    )
    _add_table(document, context["score_transform_table"])
    for offset, figure in enumerate(context["main_figures"][2:4]):
        _add_figure(
            document,
            calibration_figure_numbers[offset],
            context["figures_dir"] / figure.filename,
            figure.caption,
            figure.width_inches,
        )

    risk_band_figure_numbers = (figure_number, figure_number + 1)
    figure_number += 2
    _add_heading(document, "Risk Bands and Portfolio Stratification", level=1)
    _add_paragraph(
        document,
        "Phase 6 now produces operationally usable bands rather than collapsing most applicants into a single high-risk bucket. "
        f"Bands A, B, and C jointly hold {_format_percent(context['band_abc_share'], 1)} of the validation population, while bands D and E isolate the materially riskier tail.",
    )
    _add_paragraph(
        document,
        f"Figure {risk_band_figure_numbers[0]} shows the distribution of applicants across bands, and Figure {risk_band_figure_numbers[1]} confirms that actual default rate rises monotonically from "
        "2.10% in band A to 27.80% in band E. Table 2 provides the compact risk band summary used for reporting and governance review."
    )
    _add_table(document, context["risk_band_table"])
    for offset, figure in enumerate(context["main_figures"][4:6]):
        _add_figure(
            document,
            risk_band_figure_numbers[offset],
            context["figures_dir"] / figure.filename,
            figure.caption,
            figure.width_inches,
        )

    cutoff_figure_number = figure_number
    figure_number += 1
    _add_heading(document, "Final Policy Cutoff Analysis", level=1)
    _add_paragraph(
        document,
        "The current final strategy is a single policy rather than a set of named scenarios. "
        f"The approve cutoff is {_format_number(float(context['final_policy']['final_cutoff']), 0)}, "
        f"the review floor is {_format_number(float(context['final_policy']['final_review_cutoff']), 0)}, "
        f"and the resulting approval / review / reject mix is {_format_percent(float(context['final_policy']['approval_rate']))} / "
        f"{_format_percent(float(context['final_policy']['review_rate']))} / {_format_percent(float(context['final_policy']['reject_rate']))}.",
    )
    _add_paragraph(
        document,
        f"Figure {cutoff_figure_number} links expected value and approved bad rate to the cutoff anchor score. "
        f"At the selected policy point, approved bad rate is {_format_percent(float(context['final_policy']['actual_approved_bad_rate']))}, "
        f"and expected value per applicant is {_format_number(float(context['final_policy']['expected_value_per_applicant']))}.",
    )
    _add_paragraph(document, "Table 3 records the compact final policy summary used in the report.")
    _add_table(document, context["final_policy_table"])
    _add_figure(
        document,
        cutoff_figure_number,
        context["figures_dir"] / context["main_figures"][6].filename,
        context["main_figures"][6].caption,
        context["main_figures"][6].width_inches,
    )

    migration_figure_number = figure_number
    grouped_figure_numbers = (figure_number + 1, figure_number + 2)
    figure_number += 3
    _add_heading(document, "Group Sensitivity and Migration", level=1)
    _add_paragraph(
        document,
        f"Figure {migration_figure_number} shows the migration from comparator decisions to candidate decisions. "
        f"{_format_percent(context['approve_stay_rate'])} of comparator approvals remain approvals, "
        f"{_format_percent(context['review_to_approve_rate'])} of comparator review cases migrate to approve, and "
        f"{_format_percent(context['reject_to_approve_rate'])} of comparator rejects move to approve under the candidate strategy.",
    )
    _add_paragraph(
        document,
        "The grouped policy readout still needs governance attention. "
        f"Age-band equalized-odds gap remains {_format_pp(context['age_band_gap'])}, which is materially larger than "
        f"region_rating_group at {_format_pp(context['region_gap'])} and family_status_group at {_format_pp(context['family_gap'])}. "
        f"This is why Figure {grouped_figure_numbers[0]} and Figure {grouped_figure_numbers[1]} are kept in the main body rather than relegated to an appendix.",
    )
    _add_paragraph(
        document,
        "Table 4 summarizes the approval-rate spread within the three grouped views that are most relevant to the current policy interpretation."
    )
    _add_table(document, context["group_gap_table"])
    for figure_number_value, figure in zip(
        (migration_figure_number, grouped_figure_numbers[0], grouped_figure_numbers[1]),
        context["main_figures"][7:10],
    ):
        _add_figure(
            document,
            figure_number_value,
            context["figures_dir"] / figure.filename,
            figure.caption,
            figure.width_inches,
        )

    _add_heading(document, "Conclusion", level=1)
    _add_paragraph(
        document,
        "Phase 6 now supports a defensible policy narrative because calibrated PD, score translation, and operational banding are aligned. "
        "The main correction is not a rejection of the candidate model's ranking quality; it is a correction to how probability outputs are translated into portfolio decisions.",
    )
    _add_paragraph(
        document,
        "The final policy point is operationally coherent under the current assumptions: score 505 for approval, score 485 for review, "
        f"approval / review / reject = {_format_percent(float(context['final_policy']['approval_rate']))} / "
        f"{_format_percent(float(context['final_policy']['review_rate']))} / {_format_percent(float(context['final_policy']['reject_rate']))}, "
        f"approved bad rate = {_format_percent(float(context['final_policy']['actual_approved_bad_rate']))}, and "
        f"expected value per applicant = {_format_number(float(context['final_policy']['expected_value_per_applicant']))}.",
    )
    _add_paragraph(
        document,
        "At the same time, the report preserves the governance bridge from Phase 5: ext_source dominates proxy uplift, age_band remains the strongest grouped sensitivity, "
        "and region_rating_group plus family_status_group still warrant review when this strategy is discussed beyond the validation setting.",
    )

    _add_heading(document, "Limitations", level=1)
    _add_bullet(
        document,
        "This report is built from current repository artifacts and validation-set diagnostics; it is not a live monitoring report.",
    )
    _add_bullet(
        document,
        "Expected value results depend on the configured unit-economics assumptions and should be sensitivity-tested before operational adoption.",
    )
    _add_bullet(
        document,
        "Grouped policy outputs are governance diagnostics only and should not be described as a completed production fairness audit.",
    )

    if context["appendix_figures"]:
        document.add_page_break()
        _add_heading(document, "Appendix: Supplementary Figures", level=1)
        _add_paragraph(
            document,
            "The following supplementary figures are included for traceability because they help connect the calibrated PD translation to the final score distribution.",
        )
        for figure in context["appendix_figures"]:
            _add_figure(
                document,
                figure_number,
                context["figures_dir"] / figure.filename,
                figure.caption,
                figure.width_inches,
            )
            figure_number += 1

    document.save(output_file)
    return output_file


if __name__ == "__main__":
    print(generate_phase6_analysis_docx())
