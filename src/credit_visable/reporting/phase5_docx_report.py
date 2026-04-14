"""Generate an English Phase 5 analytical DOCX report from existing artifacts."""

from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd

from credit_visable.utils import get_paths

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - exercised only when dependency is absent.
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None
    _DOCX_IMPORT_ERROR = exc
else:
    _DOCX_IMPORT_ERROR = None


DEFAULT_REPORT_NAME = "Phase5_XAI_Fairness_Governance_Analysis_Report.docx"
PHASE5_DIR_NAME = "xai_fairness"


@dataclass(frozen=True)
class FigureSpec:
    """DOCX figure placement instructions."""

    filename: str
    caption: str
    width_inches: float


MAIN_FIGURES: tuple[FigureSpec, ...] = (
    FigureSpec(
        filename="phase5_traditional_plus_proxy_shap_beeswarm.png",
        caption="Global SHAP beeswarm for the candidate model, showing that EXT_SOURCE features dominate model-wide contribution magnitude.",
        width_inches=5.8,
    ),
    FigureSpec(
        filename="phase5_traditional_plus_proxy_shap_bar.png",
        caption="Mean absolute SHAP contribution for the candidate model, confirming EXT_SOURCE_3, EXT_SOURCE_2, and AMT_GOODS_PRICE as the main global drivers.",
        width_inches=6.0,
    ),
    FigureSpec(
        filename="phase5_proxy_family_uplift_delta.png",
        caption="Proxy family uplift delta between the candidate model and the matched traditional-core comparator, showing ext_source as the dominant uplift channel.",
        width_inches=6.1,
    ),
    FigureSpec(
        filename="phase5_fairness_metric_gaps.png",
        caption="Grouped fairness metric gaps across reviewed attributes, with age_band showing the largest disparity.",
        width_inches=6.1,
    ),
    FigureSpec(
        filename="phase5_age_band_approval_rate.png",
        caption="Approval-rate distribution by age_band, highlighting the largest grouped approval spread in the Phase 5 governance review.",
        width_inches=6.1,
    ),
    FigureSpec(
        filename="phase5_region_rating_approval_rate.png",
        caption="Approval-rate distribution by region_rating_group, showing a material disadvantage for the highest-risk regional segment.",
        width_inches=6.1,
    ),
    FigureSpec(
        filename="phase5_family_status_approval_rate.png",
        caption="Approval-rate distribution by family_status_group, showing a persistent but smaller spread than age_band and region_rating_group.",
        width_inches=6.1,
    ),
    FigureSpec(
        filename="phase5_traditional_plus_proxy_top_interactions.png",
        caption="Top SHAP interaction pairs for the candidate model, led by the EXT_SOURCE_3 and EXT_SOURCE_2 interaction.",
        width_inches=6.0,
    ),
    FigureSpec(
        filename="phase5_traditional_plus_proxy_pdp_1.png",
        caption="Partial dependence for EXT_SOURCE_3, showing that higher external score values are associated with materially lower predicted default risk.",
        width_inches=5.8,
    ),
    FigureSpec(
        filename="phase5_traditional_plus_proxy_pdp_2.png",
        caption="Partial dependence for EXT_SOURCE_2, reinforcing the same downward response pattern observed for EXT_SOURCE_3.",
        width_inches=5.8,
    ),
)

APPENDIX_FIGURES: tuple[FigureSpec, ...] = (
    FigureSpec(
        filename="phase5_traditional_plus_proxy_case_1_shap_waterfall.png",
        caption="Illustrative SHAP waterfall for the stored high-risk bad case used in the Phase 5 appendix.",
        width_inches=6.0,
    ),
    FigureSpec(
        filename="phase5_traditional_plus_proxy_case_1_lime.png",
        caption="Illustrative LIME explanation for the same stored case, included as a local surrogate cross-check rather than a replacement for SHAP.",
        width_inches=6.0,
    ),
)


def _ensure_docx_dependency() -> None:
    if Document is None or WD_ALIGN_PARAGRAPH is None or Inches is None or Pt is None:
        raise RuntimeError(
            "python-docx is required to generate the Phase 5 report. "
            "Install it with `python -m pip install python-docx`."
        ) from _DOCX_IMPORT_ERROR


def _load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def _format_percent(value: float, decimals: int = 2) -> str:
    return f"{value * 100:.{decimals}f}%"


def _format_pp(value: float, decimals: int = 2) -> str:
    return f"{value * 100:.{decimals}f} pp"


def _format_number(value: float, decimals: int = 2) -> str:
    return f"{value:,.{decimals}f}"


def _set_run_font(paragraph, *, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        run.bold = bold or run.bold
        run.italic = italic or run.italic
        if run.font is not None:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def _style_document(document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"


def _resolve_output_path(paths, output_path: Path | None) -> Path:
    if output_path is not None:
        return output_path
    return paths.reports / "Analysis" / DEFAULT_REPORT_NAME


def _build_model_metrics_table(candidate: dict[str, Any], comparator: dict[str, Any]) -> pd.DataFrame:
    rows = []
    for label, payload in [
        ("Candidate", candidate),
        ("Matched Comparator", comparator),
    ]:
        rows.append(
            {
                "Model": label,
                "Model Label": str(payload["model_label"]),
                "Feature Count": int(payload["feature_count"]),
                "ROC-AUC": f"{float(payload['roc_auc']):.4f}",
                "Average Precision": f"{float(payload['average_precision']):.4f}",
                "KS": f"{float(payload['ks_statistic']):.4f}",
            }
        )
    return pd.DataFrame(rows)


def _build_proxy_uplift_table(proxy_uplift: pd.DataFrame) -> pd.DataFrame:
    table = proxy_uplift.loc[
        :,
        [
            "proxy_family",
            "traditional_plus_proxy_mean_abs_contribution",
            "traditional_core_mean_abs_contribution",
            "mean_abs_contribution_delta_proxy_minus_core",
        ],
    ].copy()
    table.columns = [
        "Proxy Family",
        "Candidate Mean |SHAP|",
        "Comparator Mean |SHAP|",
        "Delta",
    ]
    for column in ["Candidate Mean |SHAP|", "Comparator Mean |SHAP|", "Delta"]:
        table[column] = table[column].map(lambda value: _format_number(float(value), 3))
    return table


def _build_fairness_metric_table(fairness_metric: pd.DataFrame) -> pd.DataFrame:
    focus = fairness_metric.loc[
        fairness_metric["protected_attribute"].isin(
            ["age_band", "region_rating_group", "family_status_group"]
        )
    ].copy()
    focus = focus[
        [
            "protected_attribute",
            "best_approval_group",
            "worst_approval_group",
            "demographic_parity_diff",
            "equal_opportunity_diff",
            "equalized_odds_gap",
        ]
    ]
    focus.columns = [
        "Protected Attribute",
        "Best Approval Group",
        "Worst Approval Group",
        "Demographic Parity Diff",
        "Equal Opportunity Diff",
        "Equalized Odds Gap",
    ]
    for column in [
        "Demographic Parity Diff",
        "Equal Opportunity Diff",
        "Equalized Odds Gap",
    ]:
        focus[column] = focus[column].map(lambda value: _format_pp(float(value)))
    return focus.reset_index(drop=True)


def _build_approval_spread_table(group_fairness: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, str]] = []
    for protected_attribute in ["age_band", "region_rating_group", "family_status_group"]:
        subset = group_fairness.loc[
            group_fairness["protected_attribute"] == protected_attribute
        ].copy()
        best_row = subset.loc[subset["approval_rate"].idxmax()]
        worst_row = subset.loc[subset["approval_rate"].idxmin()]
        rows.append(
            {
                "Protected Attribute": protected_attribute.replace("_", " ").title(),
                "Best Group": str(best_row["group"]),
                "Best Approval Rate": _format_percent(float(best_row["approval_rate"])),
                "Worst Group": str(worst_row["group"]),
                "Worst Approval Rate": _format_percent(float(worst_row["approval_rate"])),
                "Gap": _format_pp(float(best_row["approval_rate"] - worst_row["approval_rate"])),
                "Worst vs Best": _format_percent(float(worst_row["approval_rate_vs_best_group"])),
            }
        )
    return pd.DataFrame(rows)


def _build_report_context() -> dict[str, Any]:
    paths = get_paths()
    phase5_dir = paths.data_processed / PHASE5_DIR_NAME
    figures_dir = paths.reports_figures_v2

    source_paths = {
        "summary": phase5_dir / "summary.json",
        "selection": phase5_dir / "candidate_model_selection.json",
        "group_fairness": phase5_dir / "group_fairness_summary.csv",
        "fairness_metric": phase5_dir / "fairness_metric_summary.csv",
        "proxy_uplift": phase5_dir / "proxy_uplift_summary.csv",
        "top_interactions": phase5_dir / "top_shap_interactions.csv",
        "candidate_pdp": phase5_dir / "candidate_partial_dependence.csv",
        "global_contrib": phase5_dir / "traditional_plus_proxy" / "global_feature_contributions.csv",
        "local_shap": phase5_dir / "traditional_plus_proxy" / "local_case_explanations.csv",
        "local_lime": phase5_dir / "traditional_plus_proxy" / "lime_local_case_explanations.csv",
    }
    for spec in MAIN_FIGURES + APPENDIX_FIGURES:
        source_paths[spec.filename] = figures_dir / spec.filename

    missing = [str(path) for path in source_paths.values() if not path.exists()]
    if missing:
        raise FileNotFoundError(
            "The Phase 5 DOCX report requires existing Phase 5 artifacts. Missing files: "
            + ", ".join(missing)
        )

    summary = _load_json(source_paths["summary"])
    selection = _load_json(source_paths["selection"])
    group_fairness = _load_csv(source_paths["group_fairness"])
    fairness_metric = _load_csv(source_paths["fairness_metric"])
    proxy_uplift = _load_csv(source_paths["proxy_uplift"]).sort_values(
        "mean_abs_contribution_delta_proxy_minus_core", ascending=False
    )
    top_interactions = _load_csv(source_paths["top_interactions"])
    candidate_pdp = _load_csv(source_paths["candidate_pdp"])
    global_contrib = _load_csv(source_paths["global_contrib"])
    local_shap = _load_csv(source_paths["local_shap"])
    local_lime = _load_csv(source_paths["local_lime"])

    candidate = selection["candidate_model"]
    comparator = selection["matched_core_comparator"]
    dominant_proxy = proxy_uplift.iloc[0]

    fairness_focus = fairness_metric.set_index("protected_attribute")
    age_gap = float(fairness_focus.loc["age_band", "equalized_odds_gap"])
    region_gap = float(fairness_focus.loc["region_rating_group", "equalized_odds_gap"])
    family_gap = float(fairness_focus.loc["family_status_group", "equalized_odds_gap"])

    top_global = global_contrib.head(5).copy()
    top_features = ", ".join(top_global["raw_feature_name"].astype(str).tolist())

    top_interaction = top_interactions.iloc[0]
    second_interaction = top_interactions.iloc[1]
    third_interaction = top_interactions.iloc[2]

    pdp_summary = (
        candidate_pdp.groupby(["feature_rank", "feature"], as_index=False)
        .agg(
            min_partial_dependence=("partial_dependence", "min"),
            max_partial_dependence=("partial_dependence", "max"),
        )
        .sort_values("feature_rank")
        .reset_index(drop=True)
    )

    local_case = local_shap.loc[local_shap["case_role"] == "high_risk_bad"].copy()
    lime_case = local_lime.loc[local_lime["case_role"] == "high_risk_bad"].copy()
    local_case_meta = local_case.iloc[0]
    top_local_shap = ", ".join(local_case.head(3)["raw_feature_name"].astype(str).tolist())
    top_local_lime = ", ".join(lime_case.head(3)["raw_feature_name"].astype(str).tolist())

    return {
        "paths": paths,
        "figures_dir": figures_dir,
        "summary": summary,
        "selection": selection,
        "candidate": candidate,
        "comparator": comparator,
        "group_fairness": group_fairness,
        "fairness_metric": fairness_metric,
        "proxy_uplift": proxy_uplift,
        "top_interactions": top_interactions,
        "candidate_pdp": candidate_pdp,
        "global_contrib": global_contrib,
        "local_shap": local_shap,
        "local_lime": local_lime,
        "model_metrics_table": _build_model_metrics_table(candidate, comparator),
        "proxy_uplift_table": _build_proxy_uplift_table(proxy_uplift),
        "fairness_metric_table": _build_fairness_metric_table(fairness_metric),
        "approval_spread_table": _build_approval_spread_table(group_fairness),
        "main_figures": MAIN_FIGURES,
        "appendix_figures": APPENDIX_FIGURES,
        "roc_uplift": float(candidate["roc_auc"] - comparator["roc_auc"]),
        "ap_uplift": float(candidate["average_precision"] - comparator["average_precision"]),
        "ks_uplift": float(candidate["ks_statistic"] - comparator["ks_statistic"]),
        "dominant_proxy_family": str(dominant_proxy["proxy_family"]),
        "dominant_proxy_delta": float(
            dominant_proxy["mean_abs_contribution_delta_proxy_minus_core"]
        ),
        "top_features": top_features,
        "age_gap": age_gap,
        "region_gap": region_gap,
        "family_gap": family_gap,
        "top_interaction": top_interaction,
        "second_interaction": second_interaction,
        "third_interaction": third_interaction,
        "pdp_summary": pdp_summary,
        "local_case_meta": local_case_meta,
        "top_local_shap": top_local_shap,
        "top_local_lime": top_local_lime,
    }


def _add_heading(document, title: str, level: int) -> None:
    heading = document.add_heading(title, level=level)
    _set_run_font(heading, bold=True)


def _add_paragraph(document, text: str, *, bold_prefix: str | None = None, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        prefix_run.bold = True
        prefix_run.font.name = "Times New Roman"
        prefix_run.font.size = Pt(11)
        remainder = text[len(bold_prefix) :]
        if remainder:
            remainder_run = paragraph.add_run(remainder)
            remainder_run.italic = italic
            remainder_run.font.name = "Times New Roman"
            remainder_run.font.size = Pt(11)
    else:
        run = paragraph.add_run(text)
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def _add_bullet(document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def _add_table(document, table_frame: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(table_frame.columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, column in enumerate(table_frame.columns):
        header_cells[index].text = str(column)
        for run in header_cells[index].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    for _, row in table_frame.iterrows():
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for run in cells[index].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)


def _add_figure(document, figure_number: int, figure_path: Path, caption: str, width_inches: float) -> None:
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(figure_path), width=Inches(width_inches))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(f"Figure {figure_number}. {caption}")
    caption_run.italic = True
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(10)


def generate_phase5_analysis_docx(output_path: Path | None = None) -> Path:
    """Generate the Phase 5 English analysis report as a DOCX file."""

    _ensure_docx_dependency()
    context = _build_report_context()
    output_file = _resolve_output_path(context["paths"], output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Phase 5 XAI, Fairness, and Governance Analysis Report")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Home Credit Default Risk - English analytical report")
    subtitle_run.font.name = "Times New Roman"
    subtitle_run.font.size = Pt(12)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_text = (
        f"Prepared on {date.today().isoformat()} using existing Phase 5 artifacts. "
        f"Candidate model: {context['candidate']['model_label']}."
    )
    metadata_run = metadata.add_run(metadata_text)
    metadata_run.font.name = "Times New Roman"
    metadata_run.font.size = Pt(11)

    document.add_page_break()
    figure_number = 1

    _add_heading(document, "Executive Summary", level=1)
    _add_bullet(
        document,
        f"The candidate model is {context['candidate']['model_label']}, which outperforms the matched comparator "
        f"{context['comparator']['model_label']} on ROC-AUC ({context['candidate']['roc_auc']:.4f} vs {context['comparator']['roc_auc']:.4f}), "
        f"average precision ({context['candidate']['average_precision']:.4f} vs {context['comparator']['average_precision']:.4f}), "
        f"and KS ({context['candidate']['ks_statistic']:.4f} vs {context['comparator']['ks_statistic']:.4f}).",
    )
    _add_bullet(
        document,
        f"Global explainability is dominated by proxy-sensitive external score inputs: {context['dominant_proxy_family']} contributes the largest uplift delta "
        f"at {context['dominant_proxy_delta']:.3f}, far above every other proxy family reviewed in Phase 5.",
    )
    _add_bullet(
        document,
        "The strongest grouped governance sensitivity remains age_band, with demographic parity difference "
        f"{_format_pp(float(context['fairness_metric'].set_index('protected_attribute').loc['age_band', 'demographic_parity_diff']))} "
        f"and equalized odds gap {_format_pp(context['age_gap'])}.",
    )
    _add_bullet(
        document,
        "Phase 5 supports the candidate model as the stronger ranking engine, but it does not close governance review: "
        "proxy-sensitive uplift and grouped outcome gaps must carry forward into Phase 6 policy translation.",
    )

    _add_heading(document, "Candidate Model and Comparison Bridge", level=1)
    _add_paragraph(
        document,
        f"Phase 5 starts from the stored candidate decision rather than reopening model selection. "
        f"The candidate is {context['candidate']['model_label']} on the traditional_plus_proxy feature set, while the matched comparator is "
        f"{context['comparator']['model_label']} on traditional_core. The candidate uses {int(context['candidate']['feature_count'])} features versus "
        f"{int(context['comparator']['feature_count'])} for the comparator.",
    )
    _add_paragraph(
        document,
        f"Table 1 records the comparison bridge used throughout this report. The candidate delivers ROC-AUC uplift of {context['roc_uplift']:.4f}, "
        f"average-precision uplift of {context['ap_uplift']:.4f}, and KS uplift of {context['ks_uplift']:.4f}, which is why the later governance review focuses "
        "on whether that incremental strength is coming from defensible signals rather than whether the model ranks better in the first place."
    )
    _add_table(document, context["model_metrics_table"])

    global_figure_numbers = (figure_number, figure_number + 1, figure_number + 2)
    figure_number += 3
    _add_heading(document, "Global Explainability and Proxy Uplift", level=1)
    _add_paragraph(
        document,
        f"Figures {global_figure_numbers[0]} and {global_figure_numbers[1]} show that the candidate model is globally driven by {context['top_features']}. "
        "Three EXT_SOURCE variables sit at the top of the contribution ranking, while AMT_GOODS_PRICE and AMT_CREDIT provide the most prominent non-proxy support."
    )
    _add_paragraph(
        document,
        f"Figure {global_figure_numbers[2]} and Table 2 show that the proxy-sensitive uplift is highly concentrated rather than diffuse. "
        f"The ext_source family alone contributes mean absolute uplift delta {context['dominant_proxy_delta']:.3f}, while the next-largest proxy family, "
        f"{context['proxy_uplift'].iloc[1]['proxy_family']}, is only {float(context['proxy_uplift'].iloc[1]['mean_abs_contribution_delta_proxy_minus_core']):.3f}. "
        "That concentration matters because it narrows the governance question: the candidate's incremental strength is not spread evenly across all proxy families."
    )
    _add_paragraph(
        document,
        "Table 2 summarizes the compact proxy-family uplift review used in the narrative."
    )
    _add_table(document, context["proxy_uplift_table"])
    for offset, figure in enumerate(context["main_figures"][:3]):
        _add_figure(
            document,
            global_figure_numbers[offset],
            context["figures_dir"] / figure.filename,
            figure.caption,
            figure.width_inches,
        )

    fairness_figure_numbers = (
        figure_number,
        figure_number + 1,
        figure_number + 2,
        figure_number + 3,
    )
    figure_number += 4
    _add_heading(document, "Grouped Fairness and Governance Review", level=1)
    _add_paragraph(
        document,
        f"Figure {fairness_figure_numbers[0]} turns the grouped governance review into comparable gap metrics. "
        f"Age_band remains the largest reviewed sensitivity at {_format_pp(context['age_gap'])} equalized-odds gap, followed by "
        f"region_rating_group at {_format_pp(context['region_gap'])} and family_status_group at {_format_pp(context['family_gap'])}. "
        "These are governance diagnostics based on grouped validation outcomes, not a completed fairness certification."
    )
    _add_paragraph(
        document,
        f"Figures {fairness_figure_numbers[1]}, {fairness_figure_numbers[2]}, and {fairness_figure_numbers[3]} show the same issue in approval-rate terms. "
        "The widest spread is within age_band, where approval moves from 93.15% for [65,70) to 45.07% for [20,25). "
        "Region_rating_group ranges from 86.60% for group 1 to 55.20% for group 3, and family_status_group ranges from 81.24% for Widow to 61.40% for Civil marriage."
    )
    _add_paragraph(
        document,
        "Table 3 records the fairness metric summary, and Table 4 records the corresponding best-versus-worst approval spread summary used in the governance discussion."
    )
    _add_table(document, context["fairness_metric_table"])
    _add_table(document, context["approval_spread_table"])
    for offset, figure in enumerate(context["main_figures"][3:7]):
        _add_figure(
            document,
            fairness_figure_numbers[offset],
            context["figures_dir"] / figure.filename,
            figure.caption,
            figure.width_inches,
        )

    interaction_figure_numbers = (figure_number, figure_number + 1, figure_number + 2)
    figure_number += 3
    pdp_rank_1 = context["pdp_summary"].iloc[0]
    pdp_rank_2 = context["pdp_summary"].iloc[1]
    _add_heading(document, "Interaction and Response Diagnostics", level=1)
    _add_paragraph(
        document,
        f"Figure {interaction_figure_numbers[0]} shows that the candidate model's strongest learned interaction is "
        f"{context['top_interaction']['left_raw_feature']} with {context['top_interaction']['right_raw_feature']} at strength "
        f"{float(context['top_interaction']['interaction_strength']):.4f}. The next two interactions couple external scores with "
        f"{context['second_interaction']['right_raw_feature']} ({float(context['second_interaction']['interaction_strength']):.4f}) and "
        f"{context['third_interaction']['right_raw_feature']} ({float(context['third_interaction']['interaction_strength']):.4f}), which helps explain why the candidate model "
        "outperforms a simpler core-only comparator."
    )
    _add_paragraph(
        document,
        f"Figures {interaction_figure_numbers[1]} and {interaction_figure_numbers[2]} provide the response-shape check. "
        f"For {str(pdp_rank_1['feature']).replace('numeric__', '')}, partial dependence spans from "
        f"{_format_percent(float(pdp_rank_1['max_partial_dependence']))} down to {_format_percent(float(pdp_rank_1['min_partial_dependence']))}. "
        f"For {str(pdp_rank_2['feature']).replace('numeric__', '')}, the range is {_format_percent(float(pdp_rank_2['max_partial_dependence']))} down to "
        f"{_format_percent(float(pdp_rank_2['min_partial_dependence']))}. In both cases, stronger external signal is associated with lower predicted default risk."
    )
    for offset, figure in enumerate(context["main_figures"][7:10]):
        _add_figure(
            document,
            interaction_figure_numbers[offset],
            context["figures_dir"] / figure.filename,
            figure.caption,
            figure.width_inches,
        )

    _add_heading(document, "Conclusion", level=1)
    _add_paragraph(
        document,
        f"Phase 5 supports {context['candidate']['model_label']} as the stronger candidate for downstream score and cutoff translation. "
        f"Its improvement over {context['comparator']['model_label']} is consistent across ROC-AUC, average precision, and KS, which means the candidate is not being advanced on a marginal or noisy ranking gain."
    )
    _add_paragraph(
        document,
        f"At the same time, the uplift is not governance-neutral. Ext_source is the dominant proxy-sensitive family, age_band is the strongest grouped sensitivity, "
        f"and region_rating_group plus family_status_group remain material enough to carry into Phase 6. The correct Phase 5 conclusion is therefore conditional: "
        "the candidate model is analytically stronger, but its advantage must be translated into policy with explicit attention to proxy concentration and grouped outcome gaps."
    )

    _add_heading(document, "Limitations", level=1)
    _add_bullet(
        document,
        "This report is built from stored validation artifacts in the repository; it is not a live monitoring or post-deployment fairness report.",
    )
    _add_bullet(
        document,
        "Grouped metrics here are governance diagnostics only and should not be described as a production fairness clearance, legal determination, or causal claim.",
    )
    _add_bullet(
        document,
        "The report focuses on the stored candidate/comparator pair and does not reopen alternative model families, thresholds, or external policy constraints.",
    )

    document.add_page_break()
    appendix_figure_numbers = (figure_number, figure_number + 1)
    _add_heading(document, "Appendix: Local Explanation Example", level=1)
    _add_paragraph(
        document,
        f"The appendix uses the stored high-risk bad case with SK_ID_CURR {int(context['local_case_meta']['SK_ID_CURR'])}, candidate predicted PD "
        f"{_format_percent(float(context['local_case_meta']['candidate_predicted_pd']))}, and observed target {int(context['local_case_meta']['TARGET'])}. "
        f"SHAP highlights {context['top_local_shap']} as the main drivers, while LIME also places {context['top_local_lime']} at the top of its local surrogate explanation."
    )
    _add_paragraph(
        document,
        "These local plots are included for interpretability traceability only. They help illustrate the kind of case-level explanation available in Phase 5, "
        "but they should be read as supporting evidence rather than as a substitute for the global uplift and grouped governance review."
    )
    for offset, figure in enumerate(context["appendix_figures"]):
        _add_figure(
            document,
            appendix_figure_numbers[offset],
            context["figures_dir"] / figure.filename,
            figure.caption,
            figure.width_inches,
        )

    document.save(output_file)
    return output_file


if __name__ == "__main__":
    print(generate_phase5_analysis_docx())
