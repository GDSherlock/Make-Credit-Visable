"""Generate an English Phase 1-2 analytical DOCX report from existing artifacts."""

from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd

from credit_visable.utils import get_paths

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - exercised only when dependency is absent.
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None
    _DOCX_IMPORT_ERROR = exc
else:
    _DOCX_IMPORT_ERROR = None


DEFAULT_REPORT_NAME = "Phase1_Phase2_EDA_Preprocessing_Analysis_Report.docx"


@dataclass(frozen=True)
class FigureSpec:
    """DOCX figure placement instructions."""

    filename: str
    caption: str
    width_inches: float


MAIN_FIGURES: tuple[FigureSpec, ...] = (
    FigureSpec(
        filename="phase1_main_target_distribution.png",
        caption="Phase 1 target distribution showing the class imbalance that shapes downstream evaluation choices.",
        width_inches=5.8,
    ),
    FigureSpec(
        filename="phase1_main_missingness_top20.png",
        caption="Phase 1 top missingness view highlighting the housing-quality variables with the heaviest documentation gaps.",
        width_inches=5.9,
    ),
    FigureSpec(
        filename="phase1_main_correlation_heatmap.png",
        caption="Phase 1 numeric correlation heatmap used to identify concentration and redundancy among the main amount and score variables.",
        width_inches=5.8,
    ),
    FigureSpec(
        filename="phase1_top20_iv_features.png",
        caption="Phase 1 information value ranking confirming that EXT_SOURCE features dominate the univariate separation signal.",
        width_inches=5.9,
    ),
    FigureSpec(
        filename="phase1_age_band_default_rate.png",
        caption="Phase 1 age-band default rate view translating lifecycle variation into a business-readable risk slice.",
        width_inches=5.8,
    ),
    FigureSpec(
        filename="phase1_age_income_default_heatmap.png",
        caption="Phase 1 age-by-income interaction heatmap used to identify concentrated high-risk cells for later feature-governance discussion.",
        width_inches=6.0,
    ),
    FigureSpec(
        filename="phase1_name_family_status_default_rate_slice.png",
        caption="Phase 1 family-status default-rate slice illustrating that grouped differences are material but still descriptive at this stage.",
        width_inches=5.9,
    ),
)

APPENDIX_FIGURES: tuple[FigureSpec, ...] = (
    FigureSpec(
        filename="phase1_age_band_woe_trend.png",
        caption="Supplementary age-band WOE trend showing the directional scorecard-style risk pattern behind the main age slice.",
        width_inches=5.8,
    ),
    FigureSpec(
        filename="phase1_previous_application_name_contract_status_distribution.png",
        caption="Supplementary previous-application status distribution included to document the historical-table coverage considered in Phase 1.",
        width_inches=6.0,
    ),
)


def _ensure_docx_dependency() -> None:
    if Document is None or WD_ALIGN_PARAGRAPH is None or Inches is None or Pt is None:
        raise RuntimeError(
            "python-docx is required to generate the Phase 1-2 report. "
            "Install it with `python -m pip install python-docx`."
        ) from _DOCX_IMPORT_ERROR


def _load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)


def _format_percent(value: float, decimals: int = 2) -> str:
    return f"{value * 100:.{decimals}f}%"


def _format_number(value: float, decimals: int = 2) -> str:
    return f"{value:,.{decimals}f}"


def _format_shape(shape: list[int]) -> str:
    return f"{int(shape[0]):,} x {int(shape[1]):,}"


def _set_run_font(paragraph, *, bold: bool = False, italic: bool = False) -> None:
    for run in paragraph.runs:
        run.bold = bold or run.bold
        run.italic = italic or run.italic
        if run.font is not None:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def _style_document(document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"


def _resolve_output_path(paths, output_path: Path | None) -> Path:
    if output_path is not None:
        return output_path
    return paths.reports / "Analysis" / DEFAULT_REPORT_NAME


def _build_history_table(history_table_overview: pd.DataFrame) -> pd.DataFrame:
    history_table = history_table_overview.copy()
    history_table = history_table.loc[history_table["loaded"].astype(bool)].copy()
    history_table["Rows"] = history_table["rows"].map(lambda value: f"{int(value):,}")
    history_table["Columns"] = history_table["columns"].map(lambda value: f"{int(value)}")
    history_table["Memory (MB)"] = history_table["memory_mb"].map(lambda value: f"{float(value):.2f}")
    return history_table.loc[:, ["table_name", "Rows", "Columns", "Memory (MB)"]].rename(
        columns={"table_name": "History Table"}
    )


def _build_regime_comparison_table(
    core_manifest: dict[str, Any],
    proxy_manifest: dict[str, Any],
    core_decision_manifest: dict[str, Any],
    proxy_decision_manifest: dict[str, Any],
) -> pd.DataFrame:
    rows = [
        {
            "Feature Regime": "traditional_core",
            "Selected Features": int(core_decision_manifest["selected_feature_count"]),
            "Numeric": int(core_decision_manifest["numeric_feature_count"]),
            "Categorical": int(core_decision_manifest["categorical_feature_count"]),
            "Proxy Features": int(core_decision_manifest["proxy_feature_count"]),
            "Train Matrix": _format_shape(core_manifest["train_shape"]),
            "Valid Matrix": _format_shape(core_manifest["valid_shape"]),
            "Train Density": _format_percent(float(core_manifest["train_density"])),
            "Valid Density": _format_percent(float(core_manifest["valid_density"])),
        },
        {
            "Feature Regime": "traditional_plus_proxy",
            "Selected Features": int(proxy_decision_manifest["selected_feature_count"]),
            "Numeric": int(proxy_decision_manifest["numeric_feature_count"]),
            "Categorical": int(proxy_decision_manifest["categorical_feature_count"]),
            "Proxy Features": int(proxy_decision_manifest["proxy_feature_count"]),
            "Train Matrix": _format_shape(proxy_manifest["train_shape"]),
            "Valid Matrix": _format_shape(proxy_manifest["valid_shape"]),
            "Train Density": _format_percent(float(proxy_manifest["train_density"])),
            "Valid Density": _format_percent(float(proxy_manifest["valid_density"])),
        },
    ]
    return pd.DataFrame(rows)


def _build_preprocessing_rules_table(
    processing_methods_summary: dict[str, Any],
    core_manifest: dict[str, Any],
    core_decision_manifest: dict[str, Any],
) -> pd.DataFrame:
    validation_size = float(core_manifest["options"]["validation_size"])
    split_text = f"{int((1 - validation_size) * 100)}/{int(validation_size * 100)} stratified split"
    alignment_ok = bool(processing_methods_summary["cross_feature_set_split_alignment_ok"])
    clip_quantiles = core_manifest["options"]["clip_quantiles"]

    rows = [
        {"Rule": "Train / valid split", "Setting": split_text},
        {
            "Rule": "Cross-regime split alignment",
            "Setting": "Confirmed" if alignment_ok else "Not confirmed",
        },
        {
            "Rule": "Numeric imputation",
            "Setting": str(core_decision_manifest["numeric_imputation_strategy"]),
        },
        {
            "Rule": "Categorical imputation",
            "Setting": str(core_decision_manifest["categorical_imputation_strategy"]),
        },
        {"Rule": "Categorical encoding", "Setting": "One-hot with infrequent bucket"},
        {
            "Rule": "Rare-category threshold",
            "Setting": str(core_decision_manifest["rare_category_min_frequency"]),
        },
        {
            "Rule": "Numeric scaling",
            "Setting": "Disabled"
            if not bool(core_decision_manifest["numeric_scaling_enabled"])
            else "Enabled",
        },
        {
            "Rule": "Quantile clipping",
            "Setting": "Disabled" if clip_quantiles is None else str(clip_quantiles),
        },
        {
            "Rule": "Missing indicators",
            "Setting": "Disabled"
            if not bool(core_decision_manifest["missing_indicator_enabled"])
            else "Enabled",
        },
        {"Rule": "WOE transforms in main pipeline", "Setting": "Disabled"},
        {"Rule": "Output format", "Setting": "Sparse .npz matrices plus manifests"},
    ]
    return pd.DataFrame(rows)


def _add_heading(document, title: str, level: int) -> None:
    heading = document.add_heading(title, level=level)
    _set_run_font(heading, bold=True)


def _add_paragraph(document, text: str, *, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def _add_bullet(document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def _add_table(document, table_frame: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(table_frame.columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, column in enumerate(table_frame.columns):
        header_cells[index].text = str(column)
        for run in header_cells[index].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    for _, row in table_frame.iterrows():
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for run in cells[index].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)


def _add_figure(document, figure_number: int, figure_path: Path, caption: str, width_inches: float) -> None:
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(figure_path), width=Inches(width_inches))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(f"Figure {figure_number}. {caption}")
    caption_run.italic = True
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(10)


def _build_report_context() -> dict[str, Any]:
    paths = get_paths()
    figures_dir = paths.reports_figures_v2

    source_paths = {
        "eda_summary": paths.data_processed / "eda" / "eda_summary.json",
        "top_missingness": paths.data_processed / "eda" / "top_missingness.csv",
        "iv_summary": paths.data_processed / "eda" / "iv_summary.csv",
        "fairness_summary": paths.data_processed / "eda" / "fairness_summary.csv",
        "history_table_overview": paths.data_processed / "eda" / "history_table_overview.csv",
        "processing_methods_summary": paths.data_processed / "preprocessing" / "processing_methods_summary.json",
        "preprocessing_decision_summary": paths.data_processed / "preprocessing" / "preprocessing_decision_summary.csv",
        "core_manifest": paths.data_processed / "preprocessing" / "traditional_core" / "manifest.json",
        "core_feature_set_manifest": paths.data_processed
        / "preprocessing"
        / "traditional_core"
        / "feature_set_manifest.json",
        "core_decision_manifest": paths.data_processed
        / "preprocessing"
        / "traditional_core"
        / "preprocessing_decision_manifest.json",
        "proxy_manifest": paths.data_processed / "preprocessing" / "traditional_plus_proxy" / "manifest.json",
        "proxy_feature_set_manifest": paths.data_processed
        / "preprocessing"
        / "traditional_plus_proxy"
        / "feature_set_manifest.json",
        "proxy_decision_manifest": paths.data_processed
        / "preprocessing"
        / "traditional_plus_proxy"
        / "preprocessing_decision_manifest.json",
    }
    for spec in MAIN_FIGURES + APPENDIX_FIGURES:
        source_paths[spec.filename] = figures_dir / spec.filename

    missing = [str(path) for path in source_paths.values() if not path.exists()]
    if missing:
        raise FileNotFoundError(
            "The Phase 1-2 DOCX report requires existing Phase 1 and Phase 2 artifacts. Missing files: "
            + ", ".join(missing)
        )

    eda_summary = _load_json(source_paths["eda_summary"])
    top_missingness = _load_csv(source_paths["top_missingness"])
    iv_summary = _load_csv(source_paths["iv_summary"])
    fairness_summary = _load_csv(source_paths["fairness_summary"])
    history_table_overview = _load_csv(source_paths["history_table_overview"])
    processing_methods_summary = _load_json(source_paths["processing_methods_summary"])
    preprocessing_decision_summary = _load_csv(source_paths["preprocessing_decision_summary"])
    core_manifest = _load_json(source_paths["core_manifest"])
    core_feature_set_manifest = _load_json(source_paths["core_feature_set_manifest"])
    core_decision_manifest = _load_json(source_paths["core_decision_manifest"])
    proxy_manifest = _load_json(source_paths["proxy_manifest"])
    proxy_feature_set_manifest = _load_json(source_paths["proxy_feature_set_manifest"])
    proxy_decision_manifest = _load_json(source_paths["proxy_decision_manifest"])

    top_missing = top_missingness.head(3).copy()
    top_iv = iv_summary.head(5).copy()

    gender_slice = fairness_summary.loc[
        fairness_summary["protected_attribute"] == "CODE_GENDER"
    ].copy()
    gender_slice = gender_slice.loc[gender_slice["group"].isin(["F", "M"])]
    female_rate = float(gender_slice.loc[gender_slice["group"] == "F", "target_rate"].iloc[0])
    male_rate = float(gender_slice.loc[gender_slice["group"] == "M", "target_rate"].iloc[0])

    family_slice = fairness_summary.loc[
        fairness_summary["protected_attribute"] == "NAME_FAMILY_STATUS"
    ].copy()
    family_slice = family_slice.loc[family_slice["count"] >= 1000].copy()
    highest_family = family_slice.sort_values("target_rate", ascending=False).iloc[0]
    lowest_family = family_slice.sort_values("target_rate", ascending=True).iloc[0]

    history_table = _build_history_table(history_table_overview)
    largest_history_table = history_table_overview.sort_values("memory_mb", ascending=False).iloc[0]

    regime_comparison_table = _build_regime_comparison_table(
        core_manifest=core_manifest,
        proxy_manifest=proxy_manifest,
        core_decision_manifest=core_decision_manifest,
        proxy_decision_manifest=proxy_decision_manifest,
    )
    preprocessing_rules_table = _build_preprocessing_rules_table(
        processing_methods_summary=processing_methods_summary,
        core_manifest=core_manifest,
        core_decision_manifest=core_decision_manifest,
    )

    return {
        "paths": paths,
        "figures_dir": figures_dir,
        "eda_summary": eda_summary,
        "processing_methods_summary": processing_methods_summary,
        "preprocessing_decision_summary": preprocessing_decision_summary,
        "core_manifest": core_manifest,
        "core_feature_set_manifest": core_feature_set_manifest,
        "core_decision_manifest": core_decision_manifest,
        "proxy_manifest": proxy_manifest,
        "proxy_feature_set_manifest": proxy_feature_set_manifest,
        "proxy_decision_manifest": proxy_decision_manifest,
        "missing_columns": ", ".join(
            f"{row.column} ({row.missing_pct:.2f}%)" for row in top_missing.itertuples(index=False)
        ),
        "top_iv_feature": str(top_iv.iloc[0]["feature"]),
        "top_iv_value": float(top_iv.iloc[0]["iv"]),
        "top_iv_runner_up": str(top_iv.iloc[1]["feature"]),
        "female_rate": female_rate,
        "male_rate": male_rate,
        "highest_family_group": str(highest_family["group"]),
        "highest_family_rate": float(highest_family["target_rate"]),
        "lowest_family_group": str(lowest_family["group"]),
        "lowest_family_rate": float(lowest_family["target_rate"]),
        "history_table_count": int(history_table_overview["loaded"].astype(bool).sum()),
        "largest_history_table": str(largest_history_table["table_name"]),
        "largest_history_memory_mb": float(largest_history_table["memory_mb"]),
        "history_table": history_table,
        "regime_comparison_table": regime_comparison_table,
        "preprocessing_rules_table": preprocessing_rules_table,
        "main_figures": MAIN_FIGURES,
        "appendix_figures": APPENDIX_FIGURES,
    }


def generate_phase1_phase2_analysis_docx(output_path: Path | None = None) -> Path:
    """Generate the Phase 1-2 English analysis report as a DOCX file."""

    _ensure_docx_dependency()
    context = _build_report_context()
    output_file = _resolve_output_path(context["paths"], output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Phase 1-2 EDA and Preprocessing Analysis Report")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Home Credit Default Risk - English analytical report")
    subtitle_run.font.name = "Times New Roman"
    subtitle_run.font.size = Pt(12)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_run = metadata.add_run(
        f"Prepared on {date.today().isoformat()} using existing Phase 1 and Phase 2 artifacts."
    )
    metadata_run.font.name = "Times New Roman"
    metadata_run.font.size = Pt(11)

    document.add_page_break()
    figure_number = 1

    _add_heading(document, "Executive Summary", level=1)
    _add_bullet(
        document,
        "Phase 1 reviews 307,511 applications across 127 columns, and the observed default rate is 8.07%.",
    )
    _add_bullet(
        document,
        f"Missingness is materially concentrated in housing-quality fields, led by {context['missing_columns']}.",
    )
    _add_bullet(
        document,
        f"Information value ranking is led by {context['top_iv_feature']} (IV {context['top_iv_value']:.3f}), "
        f"with {context['top_iv_runner_up']} close behind, so the strongest Phase 1 signal is also proxy-sensitive.",
    )
    _add_bullet(
        document,
        "Phase 2 standardizes two aligned feature regimes: traditional_core keeps 77 selected features, while "
        "traditional_plus_proxy keeps 120 selected features including 43 proxy features.",
    )
    _add_bullet(
        document,
        "Both regimes use the same 80/20 aligned split with 246,008 training rows and 61,503 validation rows per regime, "
        "and both rely on median / most_frequent imputation plus one-hot encoding with a 0.01 rare-category threshold.",
    )

    portfolio_figure_numbers = (figure_number, figure_number + 1)
    figure_number += 2
    _add_heading(document, "Phase 1 Portfolio Snapshot and Data Coverage", level=1)
    _add_paragraph(
        document,
        "Phase 1 is a descriptive checkpoint rather than a policy stage. The base application table contains 307,511 rows and 127 columns, "
        "and Figure 1 confirms the expected target imbalance, which is why later phases must emphasize PR, KS, and calibration rather than accuracy alone.",
    )
    _add_paragraph(
        document,
        f"Figure {portfolio_figure_numbers[1]} shows that documentation gaps are not evenly distributed. "
        "The heaviest missingness sits in housing-condition variables, with COMMONAREA and NONLIVINGAPARTMENTS fields running near 68-70% missing, "
        "which is an immediate input-quality warning for any downstream preprocessing decision.",
    )
    for offset, figure in enumerate(context["main_figures"][:2]):
        _add_figure(
            document,
            portfolio_figure_numbers[offset],
            context["figures_dir"] / figure.filename,
            figure.caption,
            figure.width_inches,
        )

    variable_figure_numbers = (figure_number, figure_number + 1, figure_number + 2, figure_number + 3)
    figure_number += 4
    _add_heading(document, "Phase 1 Variable Quality and Risk Structure", level=1)
    _add_paragraph(
        document,
        f"Figure {variable_figure_numbers[0]} highlights where numeric redundancy is already visible, especially across amount-based variables and the externally sourced score families. "
        f"Figure {variable_figure_numbers[1]} then shows that the univariate separation signal is led by {context['top_iv_feature']}, "
        f"followed by {context['top_iv_runner_up']} and EXT_SOURCE_1, with IV values far above most remaining variables.",
    )
    _add_paragraph(
        document,
        f"Figure {variable_figure_numbers[2]} and Figure {variable_figure_numbers[3]} translate those signals into business language. "
        "Age is not interpreted causally, but the lifecycle slice and the age-by-income interaction clearly show that risk concentration is not random across applicant segments. "
        "These are still EDA-derived analytical fields rather than formal Phase 2 model inputs, so they guide feature governance rather than replace the raw application columns.",
    )
    for offset, figure in enumerate(context["main_figures"][2:6]):
        _add_figure(
            document,
            variable_figure_numbers[offset],
            context["figures_dir"] / figure.filename,
            figure.caption,
            figure.width_inches,
        )

    grouped_figure_number = figure_number
    figure_number += 1
    _add_heading(document, "Phase 1 Grouped Risk Slices and Historical Table Readout", level=1)
    _add_paragraph(
        document,
        f"Figure {grouped_figure_number} keeps the grouped readout in descriptive territory but still shows material differences. "
        f"In the stored fairness summary, male applicants have a { _format_percent(context['male_rate']) } observed default rate versus "
        f"{ _format_percent(context['female_rate']) } for female applicants. Across family status, "
        f"{context['highest_family_group']} is the highest-risk large segment at { _format_percent(context['highest_family_rate']) }, while "
        f"{context['lowest_family_group']} is the lowest-risk large segment at { _format_percent(context['lowest_family_rate']) }.",
    )
    _add_paragraph(
        document,
        f"Phase 1 also confirms that {context['history_table_count']} historical tables were available for descriptive review. "
        f"Table 1 summarizes that coverage, and {context['largest_history_table']} is the heaviest table by memory footprint at "
        f"{context['largest_history_memory_mb']:.2f} MB in the stored overview.",
    )
    _add_paragraph(document, "Table 1 summarizes the historical-table coverage reviewed in Phase 1.")
    _add_table(document, context["history_table"])
    _add_figure(
        document,
        grouped_figure_number,
        context["figures_dir"] / context["main_figures"][6].filename,
        context["main_figures"][6].caption,
        context["main_figures"][6].width_inches,
    )

    _add_heading(document, "Phase 2 Feature-Regime Design", level=1)
    _add_paragraph(
        document,
        "Phase 2 changes the reporting question from description to controlled model-input design. "
        "The split logic stays fixed, while the key design choice is whether proxy-sensitive fields are excluded or retained. "
        "traditional_core keeps 77 selected features and 0 proxy features, whereas traditional_plus_proxy keeps 120 selected features and 43 proxy features.",
    )
    _add_paragraph(
        document,
        f"That regime choice changes the encoded matrices as well as the field list. traditional_core produces a 246,008 x 118 training matrix and a 61,503 x 118 validation matrix, "
        f"while traditional_plus_proxy produces a 246,008 x 189 training matrix and a 61,503 x 189 validation matrix. "
        f"Table 2 summarizes the feature counts, matrix shapes, and densities for both regimes, including the lower density of {_format_percent(float(context['proxy_manifest']['train_density']))} in the proxy-inclusive matrix.",
    )
    _add_paragraph(document, "Table 2 summarizes the fixed Phase 2 regime comparison used by later modeling phases.")
    _add_table(document, context["regime_comparison_table"])

    _add_heading(document, "Phase 2 Preprocessing Rules and Artifact Readiness", level=1)
    _add_paragraph(
        document,
        "The preprocessing policy is intentionally conservative. Numeric variables use median imputation, categorical variables use most_frequent imputation, "
        "and categorical encoding remains one-hot with infrequent-category handling at 0.01. The split alignment check across feature regimes is marked ready in the stored Phase 2 summary.",
    )
    _add_paragraph(
        document,
        f"The core pipeline still avoids several steps that would change the modeling contract: scaling is disabled, quantile clipping is disabled, missing indicators are disabled, "
        f"and WOE transforms are not applied in the main Phase 2 pipeline. This is consistent with the project note that Phase 2 is a governance and reproducibility stage rather than a modeling-upgrade stage. "
        f"In the proxy-inclusive regime, the top documented IV list again starts with {context['proxy_decision_manifest']['top_iv_features'][0]}, reinforcing the bridge from Phase 1 EDA into later comparison phases.",
    )
    _add_paragraph(document, "Table 3 records the fixed preprocessing rules and artifact-readiness settings used in Phase 2.")
    _add_table(document, context["preprocessing_rules_table"])

    _add_heading(document, "Conclusion", level=1)
    _add_paragraph(
        document,
        "Taken together, Phase 1 and Phase 2 establish a defensible handoff from descriptive analysis to reusable model inputs. "
        "Phase 1 identifies where the strongest risk signal, heaviest missingness, and most visible group differences sit; Phase 2 then freezes those observations into two explicit feature regimes rather than letting later notebooks drift across inconsistent input definitions.",
    )
    _add_paragraph(
        document,
        f"The main practical conclusion is narrow but important: the strongest raw signal is still {context['top_iv_feature']}, "
        "the housing-related variables carry substantial missingness pressure, and the formal preprocessing pipeline remains basic by design. "
        "This means later model comparisons should be interpreted as regime and algorithm comparisons on fixed artifacts, not as evidence that the upstream data contract has already been optimized.",
    )

    _add_heading(document, "Limitations", level=1)
    _add_bullet(
        document,
        "This report is built entirely from stored repository artifacts and does not rerun notebooks or regenerate figures.",
    )
    _add_bullet(
        document,
        "Phase 1 analytical fields such as AGE_YEARS and interaction heatmaps are explanatory diagnostics, not formal Phase 2 model inputs.",
    )
    _add_bullet(
        document,
        "Phase 2 uses baseline preprocessing only: no scaling, no clipping, no missing indicators, and no WOE transform inside the main pipeline.",
    )
    _add_bullet(
        document,
        "Historical tables remain descriptive in this workflow; they are reviewed for signal and coverage, but not aggregated back into a customer-level training frame here.",
    )

    if context["appendix_figures"]:
        document.add_page_break()
        _add_heading(document, "Appendix: Supplementary Figures", level=1)
        _add_paragraph(
            document,
            "The following supplementary figures are included for traceability because they extend the Phase 1 explanatory record without changing the main Phase 2 contract.",
        )
        for figure in context["appendix_figures"]:
            _add_figure(
                document,
                figure_number,
                context["figures_dir"] / figure.filename,
                figure.caption,
                figure.width_inches,
            )
            figure_number += 1

    document.save(output_file)
    return output_file


if __name__ == "__main__":
    print(generate_phase1_phase2_analysis_docx())
