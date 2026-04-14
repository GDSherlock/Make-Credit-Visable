"""Tests for the Phase 1-2 DOCX report generator."""

from __future__ import annotations

import sys
from pathlib import Path


def _ensure_src_on_path() -> None:
    root = Path(__file__).resolve().parents[1]
    src = root / "src"
    if str(src) not in sys.path:
        sys.path.insert(0, str(src))


def _collect_docx_text(document) -> str:
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_generate_phase1_phase2_docx_report(tmp_path: Path) -> None:
    _ensure_src_on_path()

    from docx import Document

    from credit_visable.reporting.phase1_phase2_docx_report import (
        APPENDIX_FIGURES,
        MAIN_FIGURES,
        generate_phase1_phase2_analysis_docx,
    )
    from credit_visable.utils import get_paths

    paths = get_paths()

    required_sources = [
        paths.data_processed / "eda" / "eda_summary.json",
        paths.data_processed / "eda" / "top_missingness.csv",
        paths.data_processed / "eda" / "iv_summary.csv",
        paths.data_processed / "eda" / "fairness_summary.csv",
        paths.data_processed / "eda" / "history_table_overview.csv",
        paths.data_processed / "preprocessing" / "processing_methods_summary.json",
        paths.data_processed / "preprocessing" / "preprocessing_decision_summary.csv",
        paths.data_processed / "preprocessing" / "traditional_core" / "manifest.json",
        paths.data_processed / "preprocessing" / "traditional_core" / "feature_set_manifest.json",
        paths.data_processed / "preprocessing" / "traditional_core" / "preprocessing_decision_manifest.json",
        paths.data_processed / "preprocessing" / "traditional_plus_proxy" / "manifest.json",
        paths.data_processed / "preprocessing" / "traditional_plus_proxy" / "feature_set_manifest.json",
        paths.data_processed / "preprocessing" / "traditional_plus_proxy" / "preprocessing_decision_manifest.json",
    ]
    required_sources.extend(paths.reports_figures_v2 / spec.filename for spec in MAIN_FIGURES + APPENDIX_FIGURES)

    missing = [str(path) for path in required_sources if not path.exists()]
    assert not missing, f"Required Phase 1/2 report sources are missing: {missing}"

    output_path = tmp_path / "phase1_phase2_report.docx"
    result = generate_phase1_phase2_analysis_docx(output_path=output_path)

    assert result == output_path
    assert result.exists()
    assert result.stat().st_size > 0

    document = Document(result)
    text = _collect_docx_text(document)

    for heading in [
        "Executive Summary",
        "Phase 1 Portfolio Snapshot and Data Coverage",
        "Phase 1 Variable Quality and Risk Structure",
        "Phase 1 Grouped Risk Slices and Historical Table Readout",
        "Phase 2 Feature-Regime Design",
        "Phase 2 Preprocessing Rules and Artifact Readiness",
        "Conclusion",
        "Limitations",
        "Appendix: Supplementary Figures",
    ]:
        assert heading in text

    for expected_snippet in [
        "8.07%",
        "77",
        "120",
        "43",
        "246,008",
        "61,503",
        "0.01",
        "EXT_SOURCE_3",
    ]:
        assert expected_snippet in text

    assert len(document.inline_shapes) >= len(MAIN_FIGURES) + len(APPENDIX_FIGURES)
